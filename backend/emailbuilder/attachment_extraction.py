"""D4-B — safe structured extraction for Email AI Engineer attachments
(Feature 14 V4).

Pipeline stops here: upload -> validate (attachment_validation.py) ->
extract (this module) -> typed ExtractedFact list + provenance. Nothing
in this module constructs an EmailBrief, calls validate_action(), or
touches an EmailDocument's `content` — that is explicitly out of scope
for D4-B (see the D4-B completion report). This module has zero import
of ai_command.py / composition.py / any mutation path, by design — a
static test in test_attachments.py asserts that stays true.

Every extractor is read-only and bounded:
  - PDF (pdfplumber): text + tables per page, page-capped. pdfplumber
    (built on pdfminer.six) is a pure parser — it never executes
    embedded JavaScript, embedded files, or PDF actions.
  - DOCX (python-docx): paragraphs/headings/tables/hyperlinks. Word
    macros live in a separate `vbaProject.bin` part python-docx never
    opens or executes; embedded objects/relationships are never
    resolved or fetched.
  - XLSX (openpyxl, read_only=True): sheet/row/cell values, row-capped.
    Formulas are read as their literal formula text (openpyxl never
    evaluates a formula) — "kept as data", never executed.
  - TXT/CSV/Markdown: structural parsing only (csv module / a small
    line-based heading+list scanner for Markdown) — no third-party
    Markdown renderer is used or needed for D4-B's structural-facts goal.
  - Image: decode-and-verify only (already done during validation) —
    dimensions/format facts only, no OCR, no semantic labeling. That is
    D4-F, not this checkpoint.

Every extracted fact is untrusted CONTENT, never an instruction — see
attachment_untrusted_wrapper.py for the boundary applied whenever (in a
later checkpoint) this content reaches an AI/LLM prompt.
"""

import csv
import io
from dataclasses import asdict, dataclass, field

import docx
import openpyxl
import pdfplumber

# --- resource-exhaustion limits --------------------------------------------

MAX_PDF_PAGES = 30
MAX_PDF_CHARS = 200_000
MAX_PDF_TABLES = 100
MAX_DOCX_PARAGRAPHS = 2_000
MAX_DOCX_TABLES = 100
MAX_DOCX_TABLE_ROWS = 500
MAX_XLSX_SHEETS = 20
MAX_XLSX_ROWS_PER_SHEET = 2_000
MAX_XLSX_CELLS_TOTAL = 20_000
MAX_CSV_ROWS = 5_000
MAX_TEXT_CHARS = 200_000

_SCANNED_PDF_NOTICE = (
    'The PDF contains little or no machine-readable text. Visual/OCR '
    'analysis is not enabled in this ingestion stage.'
)


@dataclass(frozen=True)
class ExtractedFact:
    """One atomic extracted unit of content, with provenance.

    `kind` is a coarse structural label ('text' | 'heading' | 'paragraph'
    | 'list_item' | 'row' | 'table' | 'link' | 'image'); `value` is the
    extracted content itself (a string, or a small list/dict for
    tabular/image facts — never executable, never a formula result
    obtained by evaluation); `source` is the originating file kind;
    `locator` is a human-readable position (see attachment_extraction's
    module docstring's per-format provenance examples);
    `structural_meta` carries small, optional structural notes (e.g. a
    DOCX paragraph's style name); `confidence` is populated only where
    confidence has real meaning (none of D4-B's deterministic extractors
    produce one today — reserved for a future non-deterministic path).
    """

    kind: str
    value: object
    source: str
    locator: str
    structural_meta: dict | None = None
    confidence: float | None = None

    def to_dict(self) -> dict:
        return {key: val for key, val in asdict(self).items() if val is not None}


@dataclass
class ExtractionResult:
    status: str  # 'ready' | 'failed'
    error_message: str
    facts: list[ExtractedFact] = field(default_factory=list)
    meta: dict = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)


def _failed(message: str) -> ExtractionResult:
    return ExtractionResult(status='failed', error_message=message, facts=[], meta={}, warnings=[])


# --- TXT ---------------------------------------------------------------


def extract_text(uploaded_file) -> ExtractionResult:
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    text = raw.decode('utf-8', errors='replace')
    warnings = []
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
        warnings.append('Text was truncated at the character limit.')
    facts = [ExtractedFact(kind='text', value=text, source='txt', locator='file')]
    return ExtractionResult(status='ready', error_message='', facts=facts, meta={'char_count': len(text)}, warnings=warnings)


# --- CSV -----------------------------------------------------------------


def extract_csv(uploaded_file) -> ExtractionResult:
    uploaded_file.seek(0)
    raw = uploaded_file.read().decode('utf-8', errors='replace')
    warnings = []
    try:
        reader = csv.reader(io.StringIO(raw))
        rows = list(reader)
    except csv.Error:
        return _failed('This CSV file could not be parsed.')

    if not rows:
        return _failed('This CSV file has no rows.')

    header = rows[0]
    header_len = len(header)
    facts = [ExtractedFact(kind='row', value=header, source='csv', locator='row:1', structural_meta={'is_header': True})]

    malformed_count = 0
    processed_rows = rows[1:MAX_CSV_ROWS + 1]
    for index, row in enumerate(processed_rows, start=2):
        is_malformed = len(row) != header_len
        if is_malformed:
            malformed_count += 1
        facts.append(ExtractedFact(
            kind='row', value=row, source='csv', locator=f'row:{index}',
            structural_meta={'malformed': True} if is_malformed else None,
        ))

    if len(rows) - 1 > MAX_CSV_ROWS:
        warnings.append(f'Only the first {MAX_CSV_ROWS} data rows were processed.')
    if malformed_count:
        warnings.append(f'{malformed_count} row(s) had a different number of columns than the header.')

    meta = {'header': header, 'row_count': len(rows) - 1, 'malformed_row_count': malformed_count}
    return ExtractionResult(status='ready', error_message='', facts=facts, meta=meta, warnings=warnings)


# --- Markdown --------------------------------------------------------------


def extract_markdown(uploaded_file) -> ExtractionResult:
    """Bespoke, dependency-free structural scan — headings and list items
    only. Not a CommonMark-compliant renderer (D4-B doesn't render
    Markdown to HTML, only extracts structure), so no third-party
    Markdown package is installed for this."""
    uploaded_file.seek(0)
    raw = uploaded_file.read().decode('utf-8', errors='replace')
    warnings = []
    if len(raw) > MAX_TEXT_CHARS:
        raw = raw[:MAX_TEXT_CHARS]
        warnings.append('Markdown was truncated at the character limit.')

    facts = []
    heading_count = 0
    list_item_count = 0
    for index, line in enumerate(raw.splitlines(), start=1):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = stripped[level:].strip()
            if 1 <= level <= 6 and heading_text:
                facts.append(ExtractedFact(
                    kind='heading', value=heading_text, source='markdown',
                    locator=f'markdown:heading:{index}', structural_meta={'level': level},
                ))
                heading_count += 1
                continue
        if stripped[:2] in ('- ', '* ') or (stripped[:2].rstrip('.').isdigit() and '. ' in stripped[:4]):
            item_text = stripped.lstrip('-*0123456789. ').strip()
            if item_text:
                facts.append(ExtractedFact(kind='list_item', value=item_text, source='markdown', locator=f'markdown:list_item:{index}'))
                list_item_count += 1
                continue
        facts.append(ExtractedFact(kind='paragraph', value=stripped, source='markdown', locator=f'markdown:line:{index}'))

    meta = {'heading_count': heading_count, 'list_item_count': list_item_count}
    return ExtractionResult(status='ready', error_message='', facts=facts, meta=meta, warnings=warnings)


# --- PDF ---------------------------------------------------------------


def extract_pdf(uploaded_file) -> ExtractionResult:
    uploaded_file.seek(0)
    facts = []
    warnings = []
    total_chars = 0
    table_count = 0

    try:
        with pdfplumber.open(uploaded_file) as pdf:
            page_count = len(pdf.pages)
            pages = pdf.pages[:MAX_PDF_PAGES]
            for index, page in enumerate(pages, start=1):
                try:
                    text = (page.extract_text() or '').strip()
                except Exception:  # noqa: BLE001 - never leak pdfminer internals for one bad page
                    text = ''
                    warnings.append(f'Page {index} text could not be parsed.')
                if text:
                    remaining = MAX_PDF_CHARS - total_chars
                    if remaining <= 0:
                        break
                    if len(text) > remaining:
                        text = text[:remaining]
                        warnings.append('PDF text was truncated at the character limit.')
                    total_chars += len(text)
                    facts.append(ExtractedFact(kind='text', value=text, source='pdf', locator=f'pdf:page:{index}'))
                try:
                    tables = page.extract_tables() or []
                except Exception:  # noqa: BLE001 - never leak pdfminer internals for one bad page
                    tables = []
                    warnings.append(f"Page {index}'s tables could not be parsed.")
                for t_index, table in enumerate(tables, start=1):
                    if table_count >= MAX_PDF_TABLES:
                        break
                    facts.append(ExtractedFact(
                        kind='table', value=table, source='pdf',
                        locator=f'pdf:page:{index}:table:{t_index}', structural_meta={'rows': len(table)},
                    ))
                    table_count += 1
    except Exception:  # noqa: BLE001 - never leak pdfplumber/pdfminer internals to the client
        return _failed('This PDF could not be read. It may be corrupted or password-protected.')

    if page_count > MAX_PDF_PAGES:
        warnings.append(f'Only the first {MAX_PDF_PAGES} of {page_count} pages were processed.')
    if total_chars < 40:
        warnings.append(_SCANNED_PDF_NOTICE)

    meta = {'page_count': page_count, 'pages_processed': len(pages), 'table_count': table_count}
    return ExtractionResult(status='ready', error_message='', facts=facts, meta=meta, warnings=warnings)


# --- DOCX ----------------------------------------------------------------


def extract_docx(uploaded_file) -> ExtractionResult:
    uploaded_file.seek(0)
    try:
        document = docx.Document(uploaded_file)
    except Exception:  # noqa: BLE001 - never leak python-docx/lxml internals
        return _failed('This DOCX file could not be read. It may be corrupted or not a valid Word document.')

    facts = []
    warnings = []
    paragraph_count = 0
    link_count = 0

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph_count >= MAX_DOCX_PARAGRAPHS:
            warnings.append(f'Only the first {MAX_DOCX_PARAGRAPHS} paragraphs were processed.')
            break
        style_name = paragraph.style.name if paragraph.style is not None else ''
        kind = 'heading' if style_name.lower().startswith('heading') else 'paragraph'
        facts.append(ExtractedFact(
            kind=kind, value=text, source='docx', locator=f'docx:paragraph:{index}',
            structural_meta={'style': style_name} if style_name else None,
        ))
        paragraph_count += 1

        # Hyperlinks — best-effort only ("if safely available", per
        # spec): newer python-docx exposes paragraph.hyperlinks; older
        # versions don't, and this degrades to a silent no-op rather than
        # an error either way.
        hyperlinks = getattr(paragraph, 'hyperlinks', None)
        if hyperlinks:
            for h_index, link in enumerate(hyperlinks, start=1):
                url = getattr(link, 'address', '') or ''
                link_text = getattr(link, 'text', '') or ''
                if url:
                    facts.append(ExtractedFact(
                        kind='link', value={'text': link_text, 'url': url}, source='docx',
                        locator=f'docx:paragraph:{index}:link:{h_index}',
                    ))
                    link_count += 1

    table_count = 0
    for t_index, table in enumerate(document.tables, start=1):
        if table_count >= MAX_DOCX_TABLES:
            warnings.append('Additional tables were not processed after the table limit.')
            break
        rows_data = [[cell.text.strip() for cell in row.cells] for row in table.rows[:MAX_DOCX_TABLE_ROWS]]
        facts.append(ExtractedFact(
            kind='table', value=rows_data, source='docx', locator=f'docx:table:{t_index}',
            structural_meta={'rows': len(rows_data)},
        ))
        if len(table.rows) > MAX_DOCX_TABLE_ROWS:
            warnings.append(f'Table {t_index} was truncated after {MAX_DOCX_TABLE_ROWS} rows.')
        table_count += 1

    meta = {'paragraph_count': paragraph_count, 'table_count': table_count, 'link_count': link_count}
    return ExtractionResult(status='ready', error_message='', facts=facts, meta=meta, warnings=warnings)


# --- XLSX ----------------------------------------------------------------


def extract_xlsx(uploaded_file) -> ExtractionResult:
    uploaded_file.seek(0)
    try:
        # read_only=True streams rows rather than materializing the whole
        # sheet — the row cap below bounds actual work done, not just the
        # result size.
        #
        # data_only is left at its default, False — this is a formatting
        # choice about which of TWO ALREADY-INERT strings gets returned
        # for a formula cell, never a switch that would make openpyxl
        # execute anything: openpyxl has no formula engine and has never
        # evaluated a formula in either mode. data_only=False returns the
        # formula's literal source text (e.g. "=B2*2"); data_only=True
        # would instead return Excel's last-saved cached numeric result
        # for that formula (and None if the file was never opened in a
        # real spreadsheet app to produce one). False is chosen because
        # the literal formula text is the more honest "kept as data,
        # never executed" representation — a plain number carries no
        # visible trace that a formula produced it.
        workbook = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=False)
    except Exception:  # noqa: BLE001 - never leak openpyxl internals
        return _failed('This XLSX file could not be read. It may be corrupted or not a valid Excel workbook.')

    facts = []
    warnings = []
    total_cells = 0
    sheet_names = workbook.sheetnames
    processed_sheets = sheet_names[:MAX_XLSX_SHEETS]

    try:
        for sheet_name in processed_sheets:
            sheet = workbook[sheet_name]
            row_index = 0
            for row in sheet.iter_rows(max_row=MAX_XLSX_ROWS_PER_SHEET):
                row_index += 1
                row_values = [cell.value for cell in row]
                if total_cells + len(row_values) > MAX_XLSX_CELLS_TOTAL:
                    warnings.append(f'Sheet "{sheet_name}" was truncated at the cell limit.')
                    break
                total_cells += len(row_values)
                facts.append(ExtractedFact(
                    kind='row', value=row_values, source='xlsx',
                    locator=f'xlsx:sheet:{sheet_name}!row:{row_index}',
                    structural_meta={'is_header_guess': True} if row_index == 1 else None,
                ))
            if row_index >= MAX_XLSX_ROWS_PER_SHEET:
                warnings.append(f'Sheet "{sheet_name}" was truncated after {MAX_XLSX_ROWS_PER_SHEET} rows.')
            if total_cells >= MAX_XLSX_CELLS_TOTAL:
                break
    finally:
        workbook.close()

    if len(sheet_names) > MAX_XLSX_SHEETS:
        warnings.append(f'Only the first {MAX_XLSX_SHEETS} of {len(sheet_names)} sheets were processed.')

    meta = {'sheet_names': sheet_names, 'sheets_processed': processed_sheets, 'row_fact_count': len(facts)}
    return ExtractionResult(status='ready', error_message='', facts=facts, meta=meta, warnings=warnings)


# --- image ---------------------------------------------------------------


def extract_image(probe_meta: dict, content_type: str) -> ExtractionResult:
    """No re-decode — attachment_validation.classify_and_validate_upload
    already decoded/verified the image once and captured `probe_meta`.
    Dimensions/format only. No OCR, no semantic/region labeling — that is
    D4-F, not this checkpoint."""
    facts = [ExtractedFact(kind='image', value=dict(probe_meta), source='image', locator='image:file')]
    meta = {**probe_meta, 'content_type': content_type}
    return ExtractionResult(status='ready', error_message='', facts=facts, meta=meta, warnings=[])


# --- dispatch --------------------------------------------------------------

_TEXT_FAMILY_EXTRACTORS = {
    'text': extract_text,
    'csv': extract_csv,
    'markdown': extract_markdown,
    'pdf': extract_pdf,
    'docx': extract_docx,
    'xlsx': extract_xlsx,
}


def run_extraction(detected_type: str, uploaded_file, *, probe_meta: dict | None = None, content_type: str = '') -> ExtractionResult:
    if detected_type == 'image':
        return extract_image(probe_meta or {}, content_type)
    extractor = _TEXT_FAMILY_EXTRACTORS.get(detected_type)
    if extractor is None:
        return _failed('This file type is not supported for extraction.')
    return extractor(uploaded_file)
